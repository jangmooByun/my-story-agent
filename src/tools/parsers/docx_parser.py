"""워드 문서(.docx) 파서"""

from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional


@dataclass
class DocxDocument:
    """파싱된 워드 문서"""
    file_path: str
    title: Optional[str] = None
    content: str = ""
    raw_content: str = ""
    metadata: dict = field(default_factory=dict)
    sections: list[dict] = field(default_factory=list)


class DocxParserTool:
    """워드 문서 파서"""

    name: str = "docx_parser"
    description: str = "워드 문서(.docx) 파싱"

    def __init__(self):
        self._docx_available = None

    def _check_docx_available(self) -> bool:
        """python-docx 설치 여부 확인"""
        if self._docx_available is None:
            try:
                import docx
                self._docx_available = True
            except ImportError:
                self._docx_available = False
        return self._docx_available

    def parse(self, file_path: str) -> DocxDocument:
        """워드 문서 파싱"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        if not self._check_docx_available():
            return DocxDocument(
                file_path=str(path.absolute()),
                title=path.stem,
                content="",
                raw_content="",
                metadata={"error": "python-docx 패키지가 필요합니다. pip install python-docx"}
            )

        import docx

        doc = docx.Document(str(path))

        result = DocxDocument(
            file_path=str(path.absolute())
        )

        result.metadata = self._extract_metadata(doc)
        result.title = self._extract_title(doc, path)

        paragraphs = []
        sections = []
        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                continue

            paragraphs.append(text)

            if para.style and para.style.name.startswith('Heading'):
                level = self._get_heading_level(para.style.name)

                if current_section:
                    sections.append(current_section)

                current_section = {
                    "level": level,
                    "title": text,
                    "content": ""
                }
            elif current_section:
                current_section["content"] += text + "\n"

        if current_section:
            current_section["content"] = current_section["content"].strip()
            sections.append(current_section)

        result.content = "\n\n".join(paragraphs)
        result.raw_content = result.content
        result.sections = sections

        if doc.tables:
            result.metadata["table_count"] = len(doc.tables)
            result.metadata["tables"] = self._extract_tables(doc.tables)

        return result

    def _extract_metadata(self, doc) -> dict:
        """문서 메타데이터 추출"""
        metadata = {}

        try:
            props = doc.core_properties

            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
            if props.created:
                metadata["created"] = props.created.isoformat() if props.created else None
            if props.modified:
                metadata["modified"] = props.modified.isoformat() if props.modified else None
            if props.last_modified_by:
                metadata["last_modified_by"] = props.last_modified_by

        except Exception:
            pass

        return metadata

    def _extract_title(self, doc, path: Path) -> str:
        """제목 추출"""
        try:
            if doc.core_properties.title:
                return doc.core_properties.title
        except Exception:
            pass

        for para in doc.paragraphs:
            if para.style and para.style.name == 'Heading 1':
                return para.text.strip()

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                return text[:100]

        return path.stem

    def _get_heading_level(self, style_name: str) -> int:
        """헤딩 레벨 추출"""
        try:
            if style_name.startswith('Heading '):
                return int(style_name.split()[-1])
        except (ValueError, IndexError):
            pass
        return 1

    def _extract_tables(self, tables) -> list[dict]:
        """테이블 정보 추출"""
        table_info = []

        for i, table in enumerate(tables[:10]):
            rows = []
            for row in table.rows[:20]:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)

            table_info.append({
                "index": i,
                "row_count": len(table.rows),
                "col_count": len(table.columns),
                "preview": rows[:5]
            })

        return table_info

    def extract_text_only(self, file_path: str) -> str:
        """텍스트만 추출 (간단 버전)"""
        doc = self.parse(file_path)
        return doc.content

    def extract_with_formatting(self, file_path: str) -> str:
        """포맷팅 정보 포함 추출 (마크다운 형식)"""
        if not self._check_docx_available():
            return ""

        import docx

        path = Path(file_path)
        doc = docx.Document(str(path))

        lines = []

        for para in doc.paragraphs:
            text = para.text.strip()

            if not text:
                lines.append("")
                continue

            if para.style and para.style.name.startswith('Heading'):
                level = self._get_heading_level(para.style.name)
                lines.append(f"{'#' * level} {text}")
            else:
                lines.append(text)

        return "\n".join(lines)
